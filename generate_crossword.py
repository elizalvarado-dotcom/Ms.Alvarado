#!/usr/bin/env python3
"""Generate Polynomial Vocabulary Crossword Puzzle Word document."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

OUTPUT_PATH = "/home/user/Ms.Alvarado/public/Polynomial_Crossword.docx"

# ── Vocabulary data ──────────────────────────────────────────────────────────
VOCAB = [
    (1,  "POLYNOMIAL",  "polinomio",           10),
    (2,  "TERM",        "término",              4),
    (3,  "LIKETERMS",   "términos semejantes",  9),
    (4,  "COEFFICIENT", "coeficiente",          11),
    (5,  "CONSTANT",    "constante",            8),
    (6,  "SIMPLIFY",    "simplificar",          8),
    (7,  "ADD",         "sumar",                3),
    (8,  "VARIABLE",    "variable",             8),
    (9,  "SUBTRACT",    "restar",               8),
    (10, "EXPRESSION",  "expresión",            10),
    (11, "DEGREE",      "grado",                6),
]

# Display form of the word (with space for LIKETERMS)
DISPLAY = {
    "POLYNOMIAL":  "Polynomial",
    "TERM":        "Term",
    "LIKETERMS":   "Like Terms",
    "COEFFICIENT": "Coefficient",
    "CONSTANT":    "Constant",
    "SIMPLIFY":    "Simplify",
    "ADD":         "Add",
    "VARIABLE":    "Variable",
    "SUBTRACT":    "Subtract",
    "EXPRESSION":  "Expression",
    "DEGREE":      "Degree",
}

ACROSS_CLUES = [
    (1,  "A math expression with more than one term",            "polinomio"),
    (2,  "A single part of an expression",                       "término"),
    (3,  "Terms that have the same variable and exponent",       "términos semejantes"),
    (4,  "A number in front of a variable",                      "coeficiente"),
    (5,  "A number with no variable",                            "constante"),
    (6,  "To make an expression simpler",                        "simplificar"),
    (7,  "To combine by putting together / sumar",               "sumar"),
    (8,  "Letters used to represent numbers",                    "variable"),
    (9,  "To take away",                                         "restar"),
    (10, "A math phrase with numbers and variables",             "expresión"),
    (11, "The highest exponent in a polynomial",                 "grado"),
]

DOWN_CLUES = [
    ("D1", "Something you combine when simplifying",     "TERM"),
    ("D2", "What you do to like terms",                  "ADD"),
    ("D3", "Opposite of add",                            "SUBTRACT"),
    ("D4", 'The prefix "poly" means many — this has many', "POLYNOMIAL"),
    ("D5", "What the letter x is called",                "VARIABLE"),
]

# ── Color constants ──────────────────────────────────────────────────────────
COLOR_GREEN  = RGBColor(0x1E, 0x6B, 0x2E)   # dark green
COLOR_ORANGE = RGBColor(0xBF, 0x5C, 0x00)   # dark orange
COLOR_RED    = RGBColor(0x8B, 0x00, 0x00)   # dark red
COLOR_BLUE   = RGBColor(0x00, 0x32, 0x7A)   # dark blue

# ── Helpers ──────────────────────────────────────────────────────────────────

def set_font(run, name="Calibri", size=11, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def add_paragraph(doc, text="", style=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph(style=style)
    p.alignment = align
    if text:
        run = p.add_run(text)
    return p


def add_heading_para(doc, text, size=14, color=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    """Add a bold colored heading paragraph."""
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.bold = True
    if color:
        run.font.color.rgb = color
    return p


def add_divider(doc):
    """Add a thin horizontal rule via paragraph bottom border."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "999999")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(docx_break_type())
    return p


def docx_break_type():
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    return br  # will be used differently below


def page_break(doc):
    """Insert a page break."""
    from docx.enum.text import WD_BREAK
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def blanks(word):
    """Return blank line representation: '__ __ __ ...' (one __ per letter)."""
    return " ".join(["__"] * len(word))


def set_cell_bg(cell, hex_color):
    """Set cell background shading."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def style_table(table):
    """Apply neat borders to a table."""
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "999999")
        tblBorders.append(border)
    tblPr.append(tblBorders)


def set_col_width(table, col_idx, width_inches):
    for row in table.rows:
        row.cells[col_idx].width = Inches(width_inches)


# ── Section builders ─────────────────────────────────────────────────────────

def build_title_page(doc):
    # Title
    p = add_heading_para(doc, "Polynomial Vocabulary Crossword Puzzle",
                         size=18, align=WD_ALIGN_PARAGRAPH.CENTER)

    # Subtitle
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p2.add_run("Rompecabezas de vocabulario de polinomios")
    r.font.name = "Calibri"
    r.font.size = Pt(14)
    r.font.italic = True
    r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    doc.add_paragraph()

    # Teacher line
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("Teacher: Miss Alvarado")
    r3.font.name = "Calibri"
    r3.font.size = Pt(12)
    r3.font.bold = True

    doc.add_paragraph()

    # Name / Period / Date line
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r4 = p4.add_run("Name: _________________________  Period: _______  Date: _______________")
    r4.font.name = "Calibri"
    r4.font.size = Pt(12)

    doc.add_paragraph()
    add_divider(doc)
    doc.add_paragraph()

    # Bilingual instructions
    instructions = [
        ("Instructions (English):",
         "Match each clue to the correct vocabulary word. Fill in the blanks — "
         "one underscore per letter. Use the word bank if provided."),
        ("Instrucciones (Español):",
         "Relaciona cada pista con la palabra de vocabulario correcta. Llena los espacios en blanco — "
         "un guion por letra. Usa el banco de palabras si se proporciona."),
    ]
    for label, body in instructions:
        p = doc.add_paragraph()
        r_label = p.add_run(label + "  ")
        r_label.font.name = "Calibri"
        r_label.font.size = Pt(11)
        r_label.font.bold = True
        r_body = p.add_run(body)
        r_body.font.name = "Calibri"
        r_body.font.size = Pt(11)

    doc.add_paragraph()

    # Difficulty overview table
    overview = doc.add_table(rows=4, cols=2)
    overview.style = "Table Grid"
    style_table(overview)
    overview.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = [("Level", "Description")]
    rows_data = [
        ("⭐ Level 1 – Easy",      "Word bank with English & Spanish translations provided"),
        ("⭐⭐ Level 2 – Medium",   "Word bank in English only"),
        ("⭐⭐⭐ Level 3 – Challenge", "No word bank; includes bonus DOWN clues"),
    ]
    # Header row
    hrow = overview.rows[0]
    for i, h in enumerate(headers[0]):
        cell = hrow.cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.name = "Calibri"
        set_cell_bg(cell, "D0D0D0")

    for ri, (lvl, desc) in enumerate(rows_data, start=1):
        row = overview.rows[ri]
        row.cells[0].text = lvl
        row.cells[1].text = desc
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(11)


def build_fill_grid(doc, show_clue_hint=True):
    """Add the fill-in-the-blank grid (numbered 1–11 with blanks)."""
    for num, word, spanish, count in VOCAB:
        p = doc.add_paragraph()
        r_num = p.add_run(f"{num}.  ")
        r_num.font.name = "Calibri"
        r_num.font.size = Pt(12)
        r_num.font.bold = True
        r_blanks = p.add_run(blanks(word))
        r_blanks.font.name = "Courier New"
        r_blanks.font.size = Pt(13)
        r_len = p.add_run(f"   ({count} letters)")
        r_len.font.name = "Calibri"
        r_len.font.size = Pt(10)
        r_len.font.color.rgb = RGBColor(0x88, 0x88, 0x88)


def build_across_clues(doc):
    """Add the ACROSS clues section."""
    p = doc.add_paragraph()
    r = p.add_run("ACROSS Clues  |  Pistas HORIZONTALES")
    r.font.name = "Calibri"
    r.font.size = Pt(12)
    r.font.bold = True
    r.font.underline = True

    for num, clue_en, clue_es in ACROSS_CLUES:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Inches(0.25)
        # Override the number
        p.clear()
        r1 = p.add_run(f"{num}.  ")
        r1.font.name = "Calibri"
        r1.font.size = Pt(11)
        r1.font.bold = True
        r2 = p.add_run(clue_en)
        r2.font.name = "Calibri"
        r2.font.size = Pt(11)
        r3 = p.add_run(f"  ({clue_es})")
        r3.font.name = "Calibri"
        r3.font.size = Pt(11)
        r3.font.italic = True
        r3.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def build_down_clues(doc):
    """Add the DOWN clues section (Challenge level)."""
    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run("BONUS — DOWN Clues  |  Pistas VERTICALES (Bonus)")
    r.font.name = "Calibri"
    r.font.size = Pt(12)
    r.font.bold = True
    r.font.underline = True

    note = doc.add_paragraph()
    rn = note.add_run(
        "These bonus clues share answers with the ACROSS list above. "
        "Use them as extra hints!  |  "
        "Estas pistas comparten respuestas con la lista de HORIZONTALES. ¡Úsalas como ayuda extra!"
    )
    rn.font.name = "Calibri"
    rn.font.size = Pt(10)
    rn.font.italic = True
    rn.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    for label, clue, answer in DOWN_CLUES:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        r1 = p.add_run(f"{label}.  ")
        r1.font.name = "Calibri"
        r1.font.size = Pt(11)
        r1.font.bold = True
        r2 = p.add_run(clue)
        r2.font.name = "Calibri"
        r2.font.size = Pt(11)


def build_word_bank_bilingual(doc):
    """Level 1: word bank with English + Spanish."""
    p = doc.add_paragraph()
    r = p.add_run("Word Bank  |  Banco de Palabras")
    r.font.name = "Calibri"
    r.font.size = Pt(12)
    r.font.bold = True

    # 2-column layout inside a 4-col table: [English | Spanish | English | Spanish]
    table = doc.add_table(rows=6, cols=4)
    style_table(table)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    headers = ["English", "Español", "English", "Español"]
    hrow = table.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.name = "Calibri"
        cell.paragraphs[0].runs[0].font.size = Pt(11)
        set_cell_bg(cell, "E8E8E8")

    # 11 words → 5 rows of 2 side-by-side pairs + 1 partial last row
    pairs = [(DISPLAY[w], sp) for _, w, sp, _ in VOCAB]
    # arrange: first 6 in left two columns, last 5 in right two (fill by row)
    left_pairs = pairs[:6]
    right_pairs = pairs[6:]

    for ri in range(1, 6):
        row = table.rows[ri]
        li = ri - 1
        if li < len(left_pairs):
            row.cells[0].text = left_pairs[li][0]
            row.cells[1].text = left_pairs[li][1]
        if li < len(right_pairs):
            row.cells[2].text = right_pairs[li][0]
            row.cells[3].text = right_pairs[li][1]
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(11)


def build_word_bank_english_only(doc):
    """Level 2: word bank with English only."""
    p = doc.add_paragraph()
    r = p.add_run("Word Bank  |  Banco de Palabras")
    r.font.name = "Calibri"
    r.font.size = Pt(12)
    r.font.bold = True

    # Single row of words, wrapping naturally in a 4-col table
    words = [DISPLAY[w] for _, w, _, _ in VOCAB]
    cols = 4
    import math
    nrows = math.ceil(len(words) / cols)
    table = doc.add_table(rows=nrows + 1, cols=cols)
    style_table(table)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    hrow = table.rows[0]
    for i in range(cols):
        set_cell_bg(hrow.cells[i], "E8E8E8")
    hrow.cells[0].text = "Word Bank (English)"
    for run in hrow.cells[0].paragraphs[0].runs:
        run.font.bold = True
        run.font.name = "Calibri"
        run.font.size = Pt(11)
    # Merge header across all cols
    hrow.cells[0].merge(hrow.cells[cols - 1])

    for idx, word in enumerate(words):
        ri = (idx // cols) + 1
        ci = idx % cols
        cell = table.rows[ri].cells[ci]
        cell.text = word
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.name = "Calibri"
                run.font.size = Pt(11)


def build_answer_key(doc):
    """Teacher answer key table."""
    p = doc.add_paragraph()
    r = p.add_run("🔑 Teacher Answer Key  |  Clave de Respuestas")
    r.font.name = "Calibri"
    r.font.size = Pt(14)
    r.font.bold = True
    r.font.color.rgb = COLOR_BLUE

    doc.add_paragraph()

    table = doc.add_table(rows=len(VOCAB) + 1, cols=4)
    style_table(table)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    col_headers = ["Number", "Word (English)", "Spanish", "Letter Count"]
    hrow = table.rows[0]
    for i, h in enumerate(col_headers):
        cell = hrow.cells[i]
        cell.text = h
        set_cell_bg(cell, "002060")
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.name = "Calibri"
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Data rows
    for ri, (num, word, spanish, count) in enumerate(VOCAB, start=1):
        row = table.rows[ri]
        row.cells[0].text = str(num)
        row.cells[1].text = DISPLAY[word]
        row.cells[2].text = spanish
        row.cells[3].text = str(count)

        # Alternate row shading
        bg = "EEF2FF" if ri % 2 == 0 else "FFFFFF"
        for cell in row.cells:
            set_cell_bg(cell, bg)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(11)

    doc.add_paragraph()

    # Also print the answer line per word
    note = doc.add_paragraph()
    rn = note.add_run("Answer Summary:")
    rn.font.name = "Calibri"
    rn.font.size = Pt(12)
    rn.font.bold = True

    for num, word, spanish, count in VOCAB:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        r1 = p.add_run(f"{num}. ")
        r1.font.name = "Calibri"
        r1.font.size = Pt(11)
        r1.font.bold = True
        r2 = p.add_run(f"{DISPLAY[word]}")
        r2.font.name = "Calibri"
        r2.font.size = Pt(11)
        r2.font.color.rgb = RGBColor(0x00, 0x70, 0x00)
        r3 = p.add_run(f"  ({spanish})  —  {count} letters")
        r3.font.name = "Calibri"
        r3.font.size = Pt(11)
        r3.font.color.rgb = RGBColor(0x44, 0x44, 0x44)


# ── Main ─────────────────────────────────────────────────────────────────────

def build_section_header(doc, text, color):
    add_divider(doc)
    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(14)
    r.font.bold = True
    r.font.color.rgb = color
    doc.add_paragraph()


def main():
    doc = Document()

    # Set default document margins
    from docx.shared import Inches as In
    for section in doc.sections:
        section.top_margin = In(1)
        section.bottom_margin = In(1)
        section.left_margin = In(1.1)
        section.right_margin = In(1.1)

    # ── Page 1: Title & Intro ────────────────────────────────────────────────
    build_title_page(doc)

    # ── Page 2: Level 1 – Easy ───────────────────────────────────────────────
    page_break(doc)
    build_section_header(doc, "⭐ Level 1 – Easy  |  Nivel 1 – Fácil", COLOR_GREEN)

    p_sub = doc.add_paragraph()
    rs = p_sub.add_run(
        "Use the word bank below to fill in the blanks. Each blank = one letter. "
        "|  Usa el banco de palabras para llenar los espacios. Cada espacio = una letra."
    )
    rs.font.name = "Calibri"
    rs.font.size = Pt(11)
    rs.font.italic = True

    doc.add_paragraph()
    build_word_bank_bilingual(doc)
    doc.add_paragraph()
    build_fill_grid(doc)
    doc.add_paragraph()
    build_across_clues(doc)

    # ── Page 3: Level 2 – Medium ─────────────────────────────────────────────
    page_break(doc)
    build_section_header(doc, "⭐⭐ Level 2 – Medium  |  Nivel 2 – Intermedio", COLOR_ORANGE)

    p_sub2 = doc.add_paragraph()
    rs2 = p_sub2.add_run(
        "Use the English word bank to fill in the blanks.  "
        "|  Usa el banco de palabras en inglés para llenar los espacios."
    )
    rs2.font.name = "Calibri"
    rs2.font.size = Pt(11)
    rs2.font.italic = True

    doc.add_paragraph()
    build_word_bank_english_only(doc)
    doc.add_paragraph()
    build_fill_grid(doc)
    doc.add_paragraph()
    build_across_clues(doc)

    # ── Page 4: Level 3 – Challenge ──────────────────────────────────────────
    page_break(doc)
    build_section_header(doc, "⭐⭐⭐ Level 3 – Challenge  |  Nivel 3 – Desafío", COLOR_RED)

    p_sub3 = doc.add_paragraph()
    rs3 = p_sub3.add_run(
        "No word bank! Use the clues to figure out each answer. "
        "Bonus DOWN clues are provided as extra hints.  "
        "|  ¡Sin banco de palabras! Usa las pistas para deducir cada respuesta."
    )
    rs3.font.name = "Calibri"
    rs3.font.size = Pt(11)
    rs3.font.italic = True

    doc.add_paragraph()
    build_fill_grid(doc)
    doc.add_paragraph()
    build_across_clues(doc)
    build_down_clues(doc)

    # ── Page 5: Answer Key ───────────────────────────────────────────────────
    page_break(doc)
    add_divider(doc)
    doc.add_paragraph()
    build_answer_key(doc)

    # Save
    doc.save(OUTPUT_PATH)
    print(f"Created: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
